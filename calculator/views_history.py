from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from django.http import JsonResponse, HttpResponse
from django.db.models import Q
import json
from io import BytesIO
from xml.sax.saxutils import escape as xml_escape
from datetime import datetime
from calculator.models import PlasticMaterial

from docx import Document
from docx.shared import Pt

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.utils import get_column_letter

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle


@login_required
def calculation_history(request):
    """Main history page showing all calculations from all sections"""

    # Import all section models
    from extrusion.models import ExtrusionCalculation
    from printing.models import PrintingCalculation
    from lamination.models import LaminationCalculation
    from slitting.models import SlittingCalculation
    from bag_making.models import BagMakingCalculation
    from sales.models import SalesCalculation

    # Get calculations from all sections - handle models without material field
    all_calculations = []

    # Extrusion calculations (has material field)
    try:
        extrusion_calculations = ExtrusionCalculation.objects.filter(user=request.user).select_related(
            'material').order_by('-timestamp')
        for calc in extrusion_calculations:
            calc.section = 'extrusion'
            calc.is_recent = is_recent(calc.timestamp)
            calc.display_material = get_display_material(calc)
            calc.display_machine = get_display_machine(calc)
            calc.display_customer = get_display_customer(calc)
            calc.display_order = get_display_order(calc)
            all_calculations.append(calc)
    except Exception as e:
        print(f"Error loading extrusion calculations: {e}")

    # Printing calculations (check if has material field)
    try:
        printing_calculations = PrintingCalculation.objects.filter(user=request.user).order_by('-timestamp')
        # Check if Printing model has material field
        if hasattr(PrintingCalculation, 'material'):
            printing_calculations = printing_calculations.select_related('material')
        for calc in printing_calculations:
            calc.section = 'printing'
            calc.is_recent = is_recent(calc.timestamp)
            calc.display_material = get_display_material(calc)
            calc.display_machine = get_display_machine(calc)
            calc.display_customer = get_display_customer(calc)
            calc.display_order = get_display_order(calc)
            all_calculations.append(calc)
    except Exception as e:
        print(f"Error loading printing calculations: {e}")

    # Lamination calculations (check if has material field)
    try:
        lamination_calculations = LaminationCalculation.objects.filter(user=request.user).order_by('-timestamp')
        # Check if Lamination model has material field
        if hasattr(LaminationCalculation, 'material'):
            lamination_calculations = lamination_calculations.select_related('material')
        for calc in lamination_calculations:
            calc.section = 'lamination'
            calc.is_recent = is_recent(calc.timestamp)
            calc.display_material = get_display_material(calc)
            calc.display_machine = get_display_machine(calc)
            calc.display_customer = get_display_customer(calc)
            calc.display_order = get_display_order(calc)
            all_calculations.append(calc)
    except Exception as e:
        print(f"Error loading lamination calculations: {e}")

    # Slitting calculations (check if has material field)
    try:
        slitting_calculations = SlittingCalculation.objects.filter(user=request.user).order_by('-timestamp')
        # Check if Slitting model has material field
        if hasattr(SlittingCalculation, 'material'):
            slitting_calculations = slitting_calculations.select_related('material')
        for calc in slitting_calculations:
            calc.section = 'slitting'
            calc.is_recent = is_recent(calc.timestamp)
            calc.display_material = get_display_material(calc)
            calc.display_machine = get_display_machine(calc)
            calc.display_customer = get_display_customer(calc)
            calc.display_order = get_display_order(calc)
            all_calculations.append(calc)
    except Exception as e:
        print(f"Error loading slitting calculations: {e}")

    # Bag Making calculations (check if has material field)
    try:
        bag_making_calculations = BagMakingCalculation.objects.filter(user=request.user).order_by('-timestamp')
        # Check if BagMaking model has material field
        if hasattr(BagMakingCalculation, 'material'):
            bag_making_calculations = bag_making_calculations.select_related('material')
        for calc in bag_making_calculations:
            calc.section = 'bag_making'
            calc.is_recent = is_recent(calc.timestamp)
            calc.display_material = get_display_material(calc)
            calc.display_machine = get_display_machine(calc)
            calc.display_customer = get_display_customer(calc)
            calc.display_order = get_display_order(calc)
            all_calculations.append(calc)
    except Exception as e:
        print(f"Error loading bag making calculations: {e}")

    # Sales calculations (no material field - uses input_data)
    try:
        sales_calculations = SalesCalculation.objects.filter(user=request.user).order_by('-timestamp')
        for calc in sales_calculations:
            calc.section = 'sales'
            calc.is_recent = is_recent(calc.timestamp)
            calc.display_material = get_display_material(calc)
            calc.display_machine = get_display_machine(calc)
            calc.display_customer = get_display_customer(calc)
            calc.display_order = get_display_order(calc)
            all_calculations.append(calc)
    except Exception as e:
        print(f"Error loading sales calculations: {e}")

    # Sort by timestamp
    materials = PlasticMaterial.objects.all()
    all_calculations.sort(key=lambda x: x.timestamp, reverse=True)

    context = {
        'calculations': all_calculations,
        'total_calculations': len(all_calculations),
        'materials': materials,
    }

    return render(request, 'calculator/history.html', context)


def get_display_material(calculation):
    """
    Get material information for display in history.

    Priority order matters here: explicit multi-layer/laminate structures the
    user actually selected must win over calculation.material, since several
    models (e.g. SlittingCalculation) require a non-null material FK and fall
    back to a default material internally even when the real selection was a
    multi-layer laminate captured separately via a "layers" structure. Checking
    the FK first would silently show that fallback instead of the real layers.
    """
    if hasattr(calculation, 'input_data') and calculation.input_data:
        input_data = calculation.input_data

        # Explicit multi-layer structure (list of dicts with material_id) - used by
        # slitting's roll mass/diameter (single OR multi-layer film structure),
        # lamination weight-breakdown/adhesive-components, and sales' Laminated Cost.
        if 'layers' in input_data and input_data['layers']:
            materials = []
            for layer in input_data['layers']:
                layer_material_id = layer.get('material_id') if isinstance(layer, dict) else None
                if layer_material_id:
                    try:
                        mat = PlasticMaterial.objects.get(id=layer_material_id)
                        materials.append({'name': mat.name, 'density': mat.density})
                    except (PlasticMaterial.DoesNotExist, ValueError):
                        pass
            if materials:
                return create_laminated_material_object(materials)

        # Generic "layer_material_ids" (flat list) - used by sales' universal
        # single/laminate material toggle.
        if 'layer_material_ids' in input_data and input_data['layer_material_ids']:
            materials = []
            for layer_material_id in input_data['layer_material_ids']:
                if layer_material_id:
                    try:
                        mat = PlasticMaterial.objects.get(id=layer_material_id)
                        materials.append({'name': mat.name, 'density': mat.density})
                    except (PlasticMaterial.DoesNotExist, ValueError):
                        pass
            if materials:
                return create_laminated_material_object(materials)

        # Check for material_id in input_data (single material)
        if 'material_id' in input_data and input_data['material_id']:
            try:
                material = PlasticMaterial.objects.get(id=input_data['material_id'])
                return material
            except (PlasticMaterial.DoesNotExist, ValueError):
                pass

        # Check for dimensions_material_id - used by bag_making's Packet/Bundle
        # Weight "Calculate from Dimensions" mode, a different key name than material_id
        if 'dimensions_material_id' in input_data and input_data['dimensions_material_id']:
            try:
                material = PlasticMaterial.objects.get(id=input_data['dimensions_material_id'])
                return material
            except (PlasticMaterial.DoesNotExist, ValueError):
                pass

        # Check for material_details in input_data (for laminated calculations)
        if 'material_details' in input_data and input_data['material_details']:
            materials = input_data['material_details']
            if materials and len(materials) > 0:
                return create_laminated_material_object(materials)

        # Check for material_detail in input_data (for roll calculations)
        if 'material_detail' in input_data and input_data['material_detail']:
            material_detail = input_data['material_detail']
            if material_detail and 'id' in material_detail:
                try:
                    material = PlasticMaterial.objects.get(id=material_detail['id'])
                    return material
                except (PlasticMaterial.DoesNotExist, ValueError):
                    pass

        # Check for primary_material_id and secondary_material_id (laminated structure)
        if 'primary_material_id' in input_data and input_data['primary_material_id']:
            return create_laminated_material_from_structure(input_data)

    # For Sales laminated calculations, check the layers structure in result_data
    if hasattr(calculation, 'section') and calculation.section == 'sales':
        if hasattr(calculation, 'result_data') and calculation.result_data:
            result_data = calculation.result_data
            if 'layer_details' in result_data and result_data['layer_details']:
                materials = []
                for layer in result_data['layer_details']:
                    if 'name' in layer:
                        materials.append({'name': layer['name']})
                if materials:
                    return create_laminated_material_object(materials)

    # Last resort: the raw material FK. This is a fallback for models that
    # require it (e.g. SlittingCalculation), so it must never be checked before
    # the explicit structures above.
    if hasattr(calculation, 'material') and calculation.material:
        return calculation.material

    return None


def get_display_machine(calculation):
    """Get machine name for display, checking the model field first, then input_data."""
    if hasattr(calculation, 'machine_name') and calculation.machine_name:
        return calculation.machine_name.replace('_', ' ').title()
    if hasattr(calculation, 'input_data') and calculation.input_data:
        machine = calculation.input_data.get('machine_name')
        if machine:
            return str(machine).replace('_', ' ').title()
    return ''


def get_display_customer(calculation):
    """Get customer name for display, checking the model field first, then input_data."""
    if hasattr(calculation, 'customer_name') and calculation.customer_name:
        return calculation.customer_name
    if hasattr(calculation, 'input_data') and calculation.input_data:
        customer = calculation.input_data.get('customer_name')
        if customer:
            return customer
    return ''


def get_display_order(calculation):
    """
    Get order/job name for display. Different apps named this field differently
    (order_name vs job_name) - check both, on the model field and in input_data.
    """
    for field_name in ('order_name', 'job_name'):
        if hasattr(calculation, field_name) and getattr(calculation, field_name):
            return getattr(calculation, field_name)
    if hasattr(calculation, 'input_data') and calculation.input_data:
        for key in ('order_name', 'job_name'):
            value = calculation.input_data.get(key)
            if value:
                return value
    return ''


def create_laminated_material_object(materials):
    """Create a special material object for laminated structures"""

    class LaminatedMaterial:
        def __init__(self, materials):
            self.name = self._generate_laminated_name(materials)
            self.density = self._calculate_average_density(materials)
            self.is_laminated = True
            self.layers = materials

        def _generate_laminated_name(self, materials):
            """Generate a descriptive name for laminated material"""
            if len(materials) == 1:
                material = materials[0]
                name = material.get('name', 'Unknown')
                return f"{name} (Single Layer)"

            layer_names = []
            for material in materials:
                name = material.get('name', 'Unknown')
                # Extract base material name (remove density info if present)
                base_name = name.split('(')[0].strip()
                layer_names.append(base_name)

            # Remove duplicates while preserving order
            unique_layers = []
            for layer in layer_names:
                if layer not in unique_layers:
                    unique_layers.append(layer)

            if len(unique_layers) == 1:
                return f"{unique_layers[0]} ({len(materials)}-Layer)"
            else:
                return f"{' / '.join(unique_layers)} (Laminated)"

        def _calculate_average_density(self, materials):
            """Calculate average density for laminated material"""
            densities = []
            for material in materials:
                if 'density' in material:
                    try:
                        densities.append(float(material['density']))
                    except (ValueError, TypeError):
                        pass
                # Try to extract density from material name if available
                elif 'name' in material:
                    name = material['name']
                    # Look for density in parentheses in name
                    import re
                    match = re.search(r'\(([\d.]+)\s*g/cm³\)', name)
                    if match:
                        try:
                            densities.append(float(match.group(1)))
                        except (ValueError, TypeError):
                            pass

            if densities:
                return round(sum(densities) / len(densities), 3)
            else:
                return 0.0  # Default density if none found

    return LaminatedMaterial(materials)


def create_laminated_material_from_structure(input_data):
    """Create laminated material from structured input data"""
    materials = []

    # Check primary material
    if 'primary_material_id' in input_data and input_data['primary_material_id']:
        try:
            primary_material = PlasticMaterial.objects.get(id=input_data['primary_material_id'])
            materials.append({
                'name': primary_material.name,
                'density': primary_material.density,
                'type': 'primary'
            })
        except (PlasticMaterial.DoesNotExist, ValueError):
            pass

    # Check secondary material
    if 'secondary_material_id' in input_data and input_data['secondary_material_id']:
        try:
            secondary_material = PlasticMaterial.objects.get(id=input_data['secondary_material_id'])
            materials.append({
                'name': secondary_material.name,
                'density': secondary_material.density,
                'type': 'secondary'
            })
        except (PlasticMaterial.DoesNotExist, ValueError):
            pass

    # Check third material
    if 'third_material_id' in input_data and input_data['third_material_id']:
        try:
            third_material = PlasticMaterial.objects.get(id=input_data['third_material_id'])
            materials.append({
                'name': third_material.name,
                'density': third_material.density,
                'type': 'third'
            })
        except (PlasticMaterial.DoesNotExist, ValueError):
            pass

    # Check fourth material
    if 'fourth_material_id' in input_data and input_data['fourth_material_id']:
        try:
            fourth_material = PlasticMaterial.objects.get(id=input_data['fourth_material_id'])
            materials.append({
                'name': fourth_material.name,
                'density': fourth_material.density,
                'type': 'fourth'
            })
        except (PlasticMaterial.DoesNotExist, ValueError):
            pass

    if materials:
        return create_laminated_material_object(materials)

    return None


@login_required
def download_calculation_history(request, format_type):
    """Download calculation history in various formats"""

    # Import all section models
    from extrusion.models import ExtrusionCalculation
    from printing.models import PrintingCalculation
    from lamination.models import LaminationCalculation
    from slitting.models import SlittingCalculation
    from bag_making.models import BagMakingCalculation
    from sales.models import SalesCalculation

    all_calculations = []

    # Get calculations from each section with proper handling
    try:
        extrusion_calculations = ExtrusionCalculation.objects.filter(user=request.user)
        if hasattr(ExtrusionCalculation, 'material'):
            extrusion_calculations = extrusion_calculations.select_related('material')
        for calc in extrusion_calculations:
            calc.section = 'extrusion'
            calc.display_material = get_display_material(calc)
            calc.display_machine = get_display_machine(calc)
            calc.display_customer = get_display_customer(calc)
            calc.display_order = get_display_order(calc)
            all_calculations.append(calc)
    except Exception as e:
        print(f"Error loading extrusion calculations for export: {e}")

    try:
        printing_calculations = PrintingCalculation.objects.filter(user=request.user)
        if hasattr(PrintingCalculation, 'material'):
            printing_calculations = printing_calculations.select_related('material')
        for calc in printing_calculations:
            calc.section = 'printing'
            calc.display_material = get_display_material(calc)
            calc.display_machine = get_display_machine(calc)
            calc.display_customer = get_display_customer(calc)
            calc.display_order = get_display_order(calc)
            all_calculations.append(calc)
    except Exception as e:
        print(f"Error loading printing calculations for export: {e}")

    try:
        lamination_calculations = LaminationCalculation.objects.filter(user=request.user)
        if hasattr(LaminationCalculation, 'material'):
            lamination_calculations = lamination_calculations.select_related('material')
        for calc in lamination_calculations:
            calc.section = 'lamination'
            calc.display_material = get_display_material(calc)
            calc.display_machine = get_display_machine(calc)
            calc.display_customer = get_display_customer(calc)
            calc.display_order = get_display_order(calc)
            all_calculations.append(calc)
    except Exception as e:
        print(f"Error loading lamination calculations for export: {e}")

    try:
        slitting_calculations = SlittingCalculation.objects.filter(user=request.user)
        if hasattr(SlittingCalculation, 'material'):
            slitting_calculations = slitting_calculations.select_related('material')
        for calc in slitting_calculations:
            calc.section = 'slitting'
            calc.display_material = get_display_material(calc)
            calc.display_machine = get_display_machine(calc)
            calc.display_customer = get_display_customer(calc)
            calc.display_order = get_display_order(calc)
            all_calculations.append(calc)
    except Exception as e:
        print(f"Error loading slitting calculations for export: {e}")

    try:
        bag_making_calculations = BagMakingCalculation.objects.filter(user=request.user)
        if hasattr(BagMakingCalculation, 'material'):
            bag_making_calculations = bag_making_calculations.select_related('material')
        for calc in bag_making_calculations:
            calc.section = 'bag_making'
            calc.display_material = get_display_material(calc)
            calc.display_machine = get_display_machine(calc)
            calc.display_customer = get_display_customer(calc)
            calc.display_order = get_display_order(calc)
            all_calculations.append(calc)
    except Exception as e:
        print(f"Error loading bag making calculations for export: {e}")

    try:
        sales_calculations = SalesCalculation.objects.filter(user=request.user)
        for calc in sales_calculations:
            calc.section = 'sales'
            calc.display_material = get_display_material(calc)
            calc.display_machine = get_display_machine(calc)
            calc.display_customer = get_display_customer(calc)
            calc.display_order = get_display_order(calc)
            all_calculations.append(calc)
    except Exception as e:
        print(f"Error loading sales calculations for export: {e}")

    # Format dispatcher with Word, Excel, PDF
    if format_type == 'word':
        return download_word_history(all_calculations, request.user.username)
    elif format_type == 'excel':
        return download_excel_history(all_calculations, request.user.username)
    elif format_type == 'pdf':
        return download_pdf_history(all_calculations, request.user.username)
    else:
        return JsonResponse({'error': 'Invalid format. Supported formats: word, excel, pdf'})


def download_word_history(calculations, username):
    """Download history as a Word document"""
    doc = Document()

    doc.add_heading('Calculation History', level=0)
    meta = doc.add_paragraph()
    meta.add_run(f'User: {username}\n').bold = True
    meta.add_run(f'Exported on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}\n')
    meta.add_run(f'Total Calculations: {len(calculations)}')

    for i, calc in enumerate(calculations, 1):
        material_info = 'N/A'
        if hasattr(calc, 'display_material') and calc.display_material:
            material_info = calc.display_material.name

        doc.add_heading(f'{i}. {get_calculation_type_display(calc)}', level=1)

        table = doc.add_table(rows=0, cols=2)
        table.style = 'Light Grid Accent 1'
        rows_data = [
            ('Section', get_section_name(calc)),
            ('Material', material_info),
            ('Machine', getattr(calc, 'display_machine', '') or 'N/A'),
            ('Customer', getattr(calc, 'display_customer', '') or 'N/A'),
            ('Order/Job', getattr(calc, 'display_order', '') or 'N/A'),
            ('Timestamp', calc.timestamp.strftime('%Y-%m-%d %H:%M:%S')),
        ]
        for label, value in rows_data:
            row_cells = table.add_row().cells
            row_cells[0].text = label
            row_cells[0].paragraphs[0].runs[0].bold = True
            row_cells[1].text = str(value)

        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run('Input Parameters:').bold = True
        for key, value in calc.input_data.items():
            doc.add_paragraph(f'{key}: {value}', style='List Bullet')

        p = doc.add_paragraph()
        p.add_run('Results:').bold = True
        for key, value in calc.result_data.items():
            doc.add_paragraph(f'{key}: {value}', style='List Bullet')

        doc.add_paragraph('_' * 60)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response[
        'Content-Disposition'] = f'attachment; filename="{username}_calculations_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx"'
    return response


def download_excel_history(calculations, username):
    """Download history as an Excel workbook"""
    wb = Workbook()
    ws = wb.active
    ws.title = 'Calculation History'

    headers = ['Section', 'Calculation Type', 'Material', 'Machine', 'Customer', 'Order/Job', 'Timestamp', 'Input Data', 'Result Data']
    ws.append(headers)

    header_font = Font(bold=True, color='FFFFFF')
    header_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
    for col_num in range(1, len(headers) + 1):
        cell = ws.cell(row=1, column=col_num)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center', vertical='center')

    for calc in calculations:
        material_info = 'N/A'
        if hasattr(calc, 'display_material') and calc.display_material:
            material_info = calc.display_material.name

        ws.append([
            get_section_name(calc),
            get_calculation_type_display(calc),
            material_info,
            getattr(calc, 'display_machine', '') or 'N/A',
            getattr(calc, 'display_customer', '') or 'N/A',
            getattr(calc, 'display_order', '') or 'N/A',
            calc.timestamp.strftime('%Y-%m-%d %H:%M:%S'),
            json.dumps(calc.input_data),
            json.dumps(calc.result_data)
        ])

    widths = [14, 28, 20, 16, 20, 20, 18, 50, 50]
    for col_num, width in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(col_num)].width = width

    ws.freeze_panes = 'A2'

    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response[
        'Content-Disposition'] = f'attachment; filename="{username}_calculations_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx"'
    return response


def download_pdf_history(calculations, username):
    """Download history as a PDF document"""
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5 * inch, bottomMargin=0.5 * inch)
    styles = getSampleStyleSheet()
    elements = []

    elements.append(Paragraph('Calculation History', styles['Title']))
    elements.append(Paragraph(f'User: {xml_escape(username)}', styles['Normal']))
    elements.append(Paragraph(f'Exported on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}', styles['Normal']))
    elements.append(Paragraph(f'Total Calculations: {len(calculations)}', styles['Normal']))
    elements.append(Spacer(1, 0.3 * inch))

    heading_style = ParagraphStyle('CalcHeading', parent=styles['Heading2'], spaceBefore=12, spaceAfter=6)
    label_style = ParagraphStyle('Label', parent=styles['Normal'], fontName='Helvetica-Bold')

    for i, calc in enumerate(calculations, 1):
        material_info = 'N/A'
        if hasattr(calc, 'display_material') and calc.display_material:
            material_info = calc.display_material.name

        elements.append(Paragraph(f'{i}. {xml_escape(get_calculation_type_display(calc))}', heading_style))

        meta_table_data = [
            ['Section', get_section_name(calc)],
            ['Material', material_info],
            ['Machine', getattr(calc, 'display_machine', '') or 'N/A'],
            ['Customer', getattr(calc, 'display_customer', '') or 'N/A'],
            ['Order/Job', getattr(calc, 'display_order', '') or 'N/A'],
            ['Timestamp', calc.timestamp.strftime('%Y-%m-%d %H:%M:%S')],
        ]
        meta_table = Table(meta_table_data, colWidths=[1.5 * inch, 4.5 * inch])
        meta_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('BACKGROUND', (0, 0), (0, -1), colors.whitesmoke),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]))
        elements.append(meta_table)
        elements.append(Spacer(1, 0.1 * inch))

        elements.append(Paragraph('Input Parameters:', label_style))
        for key, value in calc.input_data.items():
            elements.append(Paragraph(f'&bull; {xml_escape(str(key))}: {xml_escape(str(value))}', styles['Normal']))

        elements.append(Paragraph('Results:', label_style))
        for key, value in calc.result_data.items():
            elements.append(Paragraph(f'&bull; {xml_escape(str(key))}: {xml_escape(str(value))}', styles['Normal']))

        elements.append(Spacer(1, 0.25 * inch))

    doc.build(elements)
    buffer.seek(0)

    response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
    response[
        'Content-Disposition'] = f'attachment; filename="{username}_calculations_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf"'
    return response


def get_section_name(calculation):
    """Get the section name from calculation object"""
    if hasattr(calculation, 'section'):
        section_map = {
            'extrusion': 'Extrusion',
            'printing': 'Printing',
            'lamination': 'Lamination',
            'slitting': 'Slitting',
            'bag_making': 'Bag Making',
            'sales': 'Sales'
        }
        return section_map.get(calculation.section, 'Unknown')

    model_name = calculation.__class__.__name__
    if 'Extrusion' in model_name:
        return 'Extrusion'
    elif 'Printing' in model_name:
        return 'Printing'
    elif 'Lamination' in model_name:
        return 'Lamination'
    elif 'Slitting' in model_name:
        return 'Slitting'
    elif 'BagMaking' in model_name:
        return 'Bag Making'
    elif 'Sales' in model_name:
        return 'Sales'
    else:
        return 'Unknown'


def get_calculation_type_display(calculation):
    """Get the display name for calculation type"""
    if hasattr(calculation, 'get_calculation_type_display'):
        return calculation.get_calculation_type_display()
    elif hasattr(calculation, 'calculation_type'):
        return calculation.calculation_type.replace('_', ' ').title()
    else:
        return 'Unknown'


def is_recent(timestamp):
    """Check if timestamp is within last 7 days"""
    from datetime import datetime, timedelta
    one_week_ago = datetime.now() - timedelta(days=7)
    return timestamp.replace(tzinfo=None) >= one_week_ago
